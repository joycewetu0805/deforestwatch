"""
Génère le mémoire académique DeforestWatch-DRC au format Word (.docx).

Document complet et formaté (page de garde, table des matières, chapitres,
bibliographie, annexes) pour la L3 Data Science — UPC / FASI. Le contenu est
pré-rempli à partir du projet réellement implémenté ; les résultats chiffrés
proviennent du rapport d'entraînement (mode démo) et sont à remplacer par les
résultats sur données réelles une fois celles-ci collectées.

Usage : python -m scripts.generate_memoir
Sortie : docs/memoire_deforestwatch.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from config.settings import (
    ANALYSIS_END_YEAR,
    ANALYSIS_START_YEAR,
    ECOSYSTEM,
    LAND_COVER_CLASSES,
)
from src.utils import synthetic
from src.utils.helpers import load_json
from src.utils.logger import get_logger

log = get_logger("generate_memoir")

GREEN = RGBColor(0x0B, 0x6E, 0x2D)
AUTHOR = "Joyce A. WETUNGANI"
YEAR = 2026


# ── Helpers de mise en forme ──────────────────────────────────────────────
def title(doc, text, size=16, color=GREEN, center=False, bold=True):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def para(doc, text, size=11, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def numbered(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = str(h)
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    return t


# ── Construction du document ──────────────────────────────────────────────
def build() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _cover(doc)
    doc.add_page_break()
    _front_matter(doc)
    doc.add_page_break()
    _introduction(doc)
    doc.add_page_break()
    _chapter1(doc)
    doc.add_page_break()
    _chapter2(doc)
    doc.add_page_break()
    _chapter3(doc)
    doc.add_page_break()
    _chapter4(doc)
    doc.add_page_break()
    _conclusion(doc)
    doc.add_page_break()
    _bibliography(doc)
    doc.add_page_break()
    _annexes(doc)
    return doc


def _cover(doc):
    for txt, sz, c in [
        ("UNIVERSITÉ PROTESTANTE AU CONGO", 14, RGBColor(0, 0, 0)),
        ("Faculté des Sciences et Technologies (FASI)", 12, RGBColor(0, 0, 0)),
        ("Licence (L3) — Data Science", 12, RGBColor(0, 0, 0)),
    ]:
        title(doc, txt, size=sz, color=c, center=True)
    doc.add_paragraph("\n\n")
    title(doc, "DEFORESTWATCH-DRC", size=22, center=True)
    title(doc,
          "Surveillance et prédiction de la déforestation de la forêt équatoriale\n"
          "du Bassin du Congo par imagerie satellite et Machine Learning",
          size=13, center=True, color=RGBColor(0x33, 0x33, 0x33))
    doc.add_paragraph("\n\n\n")
    para(doc, f"Présenté par : {AUTHOR}", justify=False)
    para(doc, "Directeur de mémoire : ……………………………………", justify=False)
    para(doc, "Zone d'étude : Province du Mai-Ndombe (Inongo), RDC", justify=False)
    doc.add_paragraph("\n\n")
    title(doc, f"Année académique {YEAR-1}–{YEAR}", size=12, center=True,
          color=RGBColor(0, 0, 0))


def _front_matter(doc):
    title(doc, "Table des matières")
    para(doc, "[Dans Word : Références → Table des matières → mettre à jour. "
              "Les titres ci-dessous utilisent les styles de titres pour génération "
              "automatique.]", size=10)
    items = [
        "Introduction générale",
        "Chapitre 1 — Revue de la littérature",
        "Chapitre 2 — Méthodologie",
        "Chapitre 3 — Résultats et discussion",
        "Chapitre 4 — Implémentation",
        "Conclusion générale et perspectives",
        "Bibliographie",
        "Annexes",
    ]
    numbered(doc, items)

    title(doc, "Liste des figures et tableaux", size=14)
    bullets(doc, [
        "Figure 1 — Localisation de la zone d'étude (Mai-Ndombe).",
        "Figure 2 — Chaîne de traitement (architecture technique).",
        "Figure 3 — Évolution de la couverture forestière 2015–2025.",
        "Figure 4 — Cartes de classification (meilleur modèle).",
        "Figure 5 — Carte de risque de déforestation future.",
        "Tableau 1 — Sources de données utilisées.",
        "Tableau 2 — Comparaison des performances des modèles.",
    ])

    title(doc, "Sigles et abréviations", size=14)
    table(doc, ["Sigle", "Signification"], [
        ["RDC", "République Démocratique du Congo"],
        ["GEE", "Google Earth Engine"],
        ["NDVI", "Normalized Difference Vegetation Index"],
        ["EVI", "Enhanced Vegetation Index"],
        ["NDWI", "Normalized Difference Water Index"],
        ["NBR", "Normalized Burn Ratio"],
        ["RF", "Random Forest"],
        ["CNN", "Convolutional Neural Network"],
        ["U-Net", "Réseau de segmentation en U"],
        ["IoU", "Intersection over Union"],
        ["API", "Application Programming Interface"],
        ["JWT", "JSON Web Token"],
        ["2FA", "Authentification à deux facteurs"],
    ])


def _introduction(doc):
    title(doc, "Introduction générale")
    para(doc, "Contexte. La République Démocratique du Congo abrite la deuxième plus "
              "grande forêt tropicale humide du monde, au cœur du Bassin du Congo. Cet "
              "écosystème équatorial joue un rôle climatique mondial majeur (puits de "
              "carbone, régulation hydrologique, biodiversité). Il subit toutefois une "
              "déforestation accélérée, principalement due à l'agriculture sur brûlis, "
              "à l'exploitation forestière et à l'expansion des villages.")
    para(doc, "Problématique. Les plateformes existantes (Global Forest Watch, etc.) "
              "fournissent des données globales mais n'offrent pas d'analyse prédictive "
              "locale adaptée au contexte congolais. Comment exploiter l'imagerie "
              "satellite et le Machine Learning pour détecter, quantifier et anticiper "
              "la déforestation de la forêt équatoriale du Mai-Ndombe ?")
    para(doc, "Objectifs. (1) Détecter les zones déforestées par classification d'images "
              "Sentinel-2 ; (2) quantifier la perte forestière annuelle "
              f"({ANALYSIS_START_YEAR}–{ANALYSIS_END_YEAR}) ; (3) comparer trois modèles "
              "(Random Forest, XGBoost, U-Net) ; (4) prédire les zones à risque ; "
              "(5) livrer une plateforme d'aide à la décision (API + dashboard).")
    para(doc, "Méthodologie. Collecte via Google Earth Engine, prétraitement (indices "
              "spectraux, masquage des nuages, mosaïques), modélisation supervisée, "
              "validation croisée, puis déploiement (FastAPI, Streamlit, frontend React).")
    para(doc, "Plan. Le mémoire s'organise en quatre chapitres : revue de la littérature, "
              "méthodologie, résultats et discussion, puis implémentation technique.")


def _chapter1(doc):
    title(doc, "Chapitre 1 — Revue de la littérature")
    title(doc, "1.1. La déforestation en RDC et dans le Bassin du Congo", size=13)
    para(doc, f"L'écosystème cible est : {ECOSYSTEM}. La littérature documente les moteurs "
              "de la déforestation équatoriale (agriculture itinérante, charbon de bois, "
              "concessions) et l'importance du suivi spatialisé pour la conservation.")
    title(doc, "1.2. Télédétection et indices de végétation", size=13)
    para(doc, "Les capteurs multispectraux (Sentinel-2, Landsat) permettent de "
              "caractériser la végétation via des indices comme le NDVI, l'EVI, le NDWI "
              "et le NBR. La couverture nuageuse persistante en zone équatoriale impose "
              "des composites temporels (médianes de saison sèche).")
    title(doc, "1.3. Machine Learning pour la classification d'images satellites", size=13)
    para(doc, "Les approches pixel (Random Forest, XGBoost) restent des références "
              "robustes et interprétables. Les approches profondes (CNN) capturent le "
              "contexte spatial.")
    title(doc, "1.4. Segmentation sémantique et U-Net", size=13)
    para(doc, "L'architecture U-Net, encodeur-décodeur à connexions de saut, est l'état "
              "de l'art pour la segmentation d'images, y compris en télédétection, et "
              "sert ici à la cartographie pixel-à-pixel de la couverture du sol.")


def _chapter2(doc):
    title(doc, "Chapitre 2 — Méthodologie")
    title(doc, "2.1. Zone d'étude", size=13)
    para(doc, "Province du Mai-Ndombe, autour d'Inongo (-1,95° S ; 18,27° E), zone de "
              "~50 × 50 km, l'un des fronts de déforestation les plus actifs de la forêt "
              "équatoriale congolaise.")
    title(doc, "2.2. Données", size=13)
    table(doc, ["Source", "Description", "Résolution"], [
        ["Sentinel-2 (ESA)", "Images multispectrales (6 bandes)", "10 m"],
        ["Landsat 8/9 (NASA)", "Images multispectrales", "30 m"],
        ["Hansen GFC", "Perte forestière annuelle (vérité terrain)", "30 m"],
        ["SRTM", "Altitude, pente, aspect", "30 m"],
        ["OpenWeatherMap", "Précipitations, température", "—"],
    ])
    title(doc, "2.3. Prétraitement", size=13)
    bullets(doc, [
        "Masquage des nuages/ombres via la bande SCL de Sentinel-2.",
        "Composites médians de saison sèche (juin–septembre).",
        "Calcul vectorisé des indices NDVI, EVI, NDWI, NBR.",
        "Extraction de features : 6 bandes + 4 indices + 3 variables topographiques "
        "= 13 features par pixel ; tuiles 128×128 pour le U-Net.",
        "Split spatial (blocs) train/val/test 70/15/15 pour éviter la fuite géographique.",
    ])
    title(doc, "2.4. Classes de couverture du sol", size=13)
    table(doc, ["Code", "Classe"], [[k, v] for k, v in LAND_COVER_CLASSES.items()])
    title(doc, "2.5. Modèles et évaluation", size=13)
    para(doc, "Trois modèles sont comparés : Random Forest et XGBoost (pixel, 13 features) "
              "et U-Net (tuiles 128×128×6). Métriques : Accuracy, Precision, Recall, "
              "F1-score, AUC-ROC, IoU par classe et Mean IoU, matrice de confusion. "
              "Un modèle de risque (gradient boosting) prédit en outre les zones à "
              "déforester à court terme.")


def _chapter3(doc):
    title(doc, "Chapitre 3 — Résultats et discussion")
    para(doc, "Note : les chiffres ci-dessous proviennent de l'exécution en mode "
              "démonstration (données synthétiques réalistes) ; ils valident la chaîne "
              "de bout en bout et seront remplacés par les résultats sur données réelles "
              "Sentinel-2 une fois la collecte effectuée.")

    title(doc, "3.1. Dynamique de la couverture forestière", size=13)
    stats = synthetic.yearly_statistics()
    rows = [[s["year"], f"{s['total_forest_ha']:,.0f}", f"{s['forest_loss_ha']:,.0f}",
             f"{s['deforestation_rate']:.2f}"] for s in stats]
    table(doc, ["Année", "Forêt (ha)", "Perte (ha)", "Taux (%)"], rows)

    title(doc, "3.2. Comparaison des modèles", size=13)
    metrics_path = Path("data/processed/model_metrics.json")
    if metrics_path.exists():
        report = load_json(metrics_path)
        comp = report.get("comparison", [])
        if comp:
            headers = ["Modèle", "Accuracy", "F1-macro", "Mean IoU"]
            rows = [[c.get("Modèle"), c.get("Accuracy"), c.get("F1-macro"),
                     c.get("Mean IoU")] for c in comp]
            table(doc, headers, rows)
            para(doc, f"Meilleur modèle (F1-macro) : {report.get('best_model')}.")
    else:
        para(doc, "[Lancez `make train` puis régénérez le mémoire pour insérer le "
                  "tableau comparatif des modèles.]")

    title(doc, "3.3. Discussion", size=13)
    para(doc, "Les modèles pixel (Random Forest) offrent un bon compromis "
              "performance/interprétabilité ; les indices NDVI et la bande NIR (B8) "
              "ressortent comme les plus discriminants. Le U-Net exploite le contexte "
              "spatial. La carte de risque identifie les fronts forestiers menacés, "
              "en cohérence avec la proximité aux routes et villages.")


def _chapter4(doc):
    title(doc, "Chapitre 4 — Implémentation")
    title(doc, "4.1. Architecture logicielle", size=13)
    para(doc, "Le projet suit une architecture modulaire : collecte (src/data), "
              "prétraitement (src/preprocessing), modélisation (src/models), "
              "visualisation (src/visualization), API (src/api) et tableaux de bord.")
    title(doc, "4.2. Stack technique", size=13)
    table(doc, ["Composant", "Technologie"], [
        ["Langage", "Python 3.11"],
        ["Collecte", "Google Earth Engine, rasterio"],
        ["Big Data", "PySpark"],
        ["ML", "scikit-learn, XGBoost"],
        ["Deep Learning", "TensorFlow / Keras (U-Net)"],
        ["API", "FastAPI (JWT + 2FA)"],
        ["Dashboards", "Streamlit + frontend React (Vite/Tailwind)"],
        ["Base de données", "PostgreSQL / Supabase (repli SQLite)"],
        ["Conteneurisation", "Docker, docker-compose"],
        ["Tests", "pytest (28 tests)"],
    ])
    title(doc, "4.3. Ingestion de données réelles", size=13)
    para(doc, "Une couche d'abstraction (DataSource) permet de basculer du mode "
              "démonstration (données synthétiques) aux vraies images : il suffit de "
              "déposer les GeoTIFF dans data/raw/ et de mettre DEMO_MODE=false. Aucune "
              "autre modification de code n'est requise.")
    title(doc, "4.4. Authentification et sécurité", size=13)
    para(doc, "Hachage bcrypt des mots de passe, jetons JWT (access + refresh), "
              "OTP 2FA compatible Google Authenticator, contrôle d'accès par rôle "
              "(utilisateur/administrateur).")
    title(doc, "4.5. Tests et reproductibilité", size=13)
    para(doc, "La suite pytest couvre le prétraitement, les modèles, l'API "
              "(authentification, 2FA, routes protégées) et la couche de données. "
              "Le projet est conteneurisé et lançable via `make`.")


def _conclusion(doc):
    title(doc, "Conclusion générale et perspectives")
    para(doc, "Ce travail a abouti à une plateforme complète et fonctionnelle de "
              "surveillance de la déforestation de la forêt équatoriale du Mai-Ndombe, "
              "couvrant l'ensemble de la chaîne de valeur Data Science : de la collecte "
              "satellite jusqu'au déploiement d'un tableau de bord d'aide à la décision.")
    title(doc, "Limites", size=13)
    bullets(doc, [
        "Couverture nuageuse équatoriale limitant la disponibilité des images.",
        "Besoin de données de vérité terrain de qualité pour l'apprentissage supervisé.",
        "Résultats de démonstration à confirmer sur données réelles.",
    ])
    title(doc, "Perspectives", size=13)
    bullets(doc, [
        "Intégration de l'imagerie radar Sentinel-1 (insensible aux nuages).",
        "Modèles spatio-temporels (ConvLSTM, transformers) pour la prédiction.",
        "Système d'alertes en temps quasi réel pour les gestionnaires forestiers.",
        "Extension à d'autres provinces de la forêt équatoriale congolaise.",
    ])


def _bibliography(doc):
    title(doc, "Bibliographie")
    refs = [
        "Hansen, M. C., et al. (2013). High-Resolution Global Maps of 21st-Century "
        "Forest Cover Change. Science, 342(6160), 850–853.",
        "Ronneberger, O., Fischer, P., & Brox, T. (2015). U-Net: Convolutional Networks "
        "for Biomedical Image Segmentation. MICCAI.",
        "Breiman, L. (2001). Random Forests. Machine Learning, 45(1), 5–32.",
        "Chen, T., & Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System. KDD.",
        "Gorelick, N., et al. (2017). Google Earth Engine: Planetary-scale geospatial "
        "analysis for everyone. Remote Sensing of Environment, 202, 18–27.",
        "Drusch, M., et al. (2012). Sentinel-2: ESA's Optical High-Resolution Mission "
        "for GMES Operational Services. Remote Sensing of Environment, 120, 25–36.",
        "Tucker, C. J. (1979). Red and photographic infrared linear combinations for "
        "monitoring vegetation. Remote Sensing of Environment, 8(2), 127–150.",
        "FAO (2020). Global Forest Resources Assessment 2020. Rome.",
    ]
    for r in refs:
        doc.add_paragraph(r, style="List Number")


def _annexes(doc):
    title(doc, "Annexes")
    title(doc, "Annexe A — Structure du dépôt de code", size=13)
    para(doc, "config/, src/{data,preprocessing,models,visualization,api,utils}, "
              "streamlit_app/, frontend/, notebooks/, scripts/, tests/, data/.")
    title(doc, "Annexe B — Lancement du projet", size=13)
    bullets(doc, [
        "make seed — génère les données de démo et entraîne les modèles.",
        "make api — API FastAPI (http://localhost:8000/docs).",
        "make dashboard — tableau de bord Streamlit (http://localhost:8501).",
        "make frontend — frontend React (http://localhost:5173).",
        "make test — suite de tests pytest.",
    ])
    title(doc, "Annexe C — Endpoints principaux de l'API", size=13)
    bullets(doc, [
        "POST /api/v1/auth/register, /auth/login, /auth/verify-otp",
        "GET /api/v1/statistics, /predictions/{year}, /models, /source",
        "GET /api/v1/admin/users, /admin/logs (réservé administrateur)",
    ])
    title(doc, "Annexe D — Captures d'écran", size=13)
    para(doc, "[Insérer ici les captures du dashboard, de la documentation Swagger et "
              "du frontend React.]")


def main() -> None:
    out = Path("docs/memoire_deforestwatch.docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    doc = build()
    doc.save(out)
    log.info(f"Mémoire généré → {out} ({len(doc.paragraphs)} paragraphes)")


if __name__ == "__main__":
    main()
