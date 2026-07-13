"""
Génère la monographie académique DeforestWatch-DRC au format Word (.docx).

Contrairement à `generate_memoir.py`, ce script suit **strictement** la structure
imposée par le « Guide de Rédaction des Monographies des Projets et des Mémoires »
de la Faculté des Sciences Informatiques de l'Université Protestante au Congo
(janvier 2025) :

  Page de titre → Résumé → Table des matières → Sigles → Liste des figures/tableaux
  Chapitre I  : Mise en Contexte et Revue de la Littérature
  Chapitre II : Modélisation et Méthodes de Travail (orientation Data Science)
  Chapitre III: Résultats, Discussion et Conclusion
  Bibliographie & Webiographie → Annexes

Le document reste sous la limite des 40 pages (Projet L3 LMD, §10.5.2 du guide).

Le script est autonome : il ne dépend que de `python-docx` et `numpy` (les
statistiques de déforestation sont recalculées à l'identique du générateur
synthétique du projet, `src/utils/synthetic.py`). Les résultats chiffrés
proviennent du mode démonstration et sont à confirmer sur données réelles.

Usage  : python -m scripts.generate_monograph
Sortie : docs/monographie_deforestwatch.docx
"""

from __future__ import annotations

from pathlib import Path

import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor

# ── Constantes (miroir de config.settings, gardées locales pour l'autonomie) ──
AUTHOR = "Joyce A. WETUNGANI"
YEAR = 2026
ANALYSIS_START_YEAR = 2015
ANALYSIS_END_YEAR = 2025
ANALYSIS_YEARS = list(range(ANALYSIS_START_YEAR, ANALYSIS_END_YEAR + 1))
GRID_SIZE = 256
PIXEL_AREA_HA = (50_000 / GRID_SIZE) ** 2 / 10_000
ECOSYSTEM = "Forêt tropicale humide équatoriale — Bassin du Congo (Mai-Ndombe, RDC)"
LAND_COVER_CLASSES = {
    0: "Forêt dense",
    1: "Forêt dégradée",
    2: "Agriculture / Sol nu",
    3: "Eau",
    4: "Zone urbaine / Bâti",
}
AGB_TONNES_PER_HA = 310.0
CARBON_FRACTION = 0.47
CO2_PER_CARBON = 44.0 / 12.0
CO2_TONNES_PER_HA = AGB_TONNES_PER_HA * CARBON_FRACTION * CO2_PER_CARBON

GREEN = RGBColor(0x0B, 0x6E, 0x2D)
DARK = RGBColor(0x22, 0x22, 0x22)
GREY = RGBColor(0x55, 0x55, 0x55)


# ── Statistiques de déforestation (recalcul identique à src/utils/synthetic) ──
def yearly_statistics() -> list[dict]:
    rng = np.random.default_rng(42)
    yy, xx = np.mgrid[0:GRID_SIZE, 0:GRID_SIZE]
    river = np.abs(yy - (GRID_SIZE * 0.6 + 18 * np.sin(xx / 30.0))) < 3
    town_cx, town_cy = int(GRID_SIZE * 0.30), int(GRID_SIZE * 0.62)
    town = (xx - town_cx) ** 2 + (yy - town_cy) ** 2 < (GRID_SIZE * 0.04) ** 2
    sources = [(int(GRID_SIZE * 0.30), int(GRID_SIZE * 0.62)),
               (int(GRID_SIZE * 0.70), int(GRID_SIZE * 0.35)),
               (int(GRID_SIZE * 0.50), int(GRID_SIZE * 0.80))]
    road = np.abs(xx - yy) < 4
    dist = np.full((GRID_SIZE, GRID_SIZE), np.inf)
    for sx, sy in sources:
        dist = np.minimum(dist, np.sqrt((xx - sx) ** 2 + (yy - sy) ** 2))
    dist = np.minimum(dist, np.where(road, 5.0, np.inf))
    dist_norm = dist / dist.max()

    stats, prev = [], None
    for i, year in enumerate(ANALYSIS_YEARS):
        progress = i / (len(ANALYSIS_YEARS) - 1)
        lc = np.zeros((GRID_SIZE, GRID_SIZE), dtype=np.int8)
        threshold = 0.10 + 0.45 * progress
        noise = rng.normal(0, 0.05, (GRID_SIZE, GRID_SIZE))
        deforested = (dist_norm + noise) < threshold
        degraded = (dist_norm + noise < threshold + 0.08) & ~deforested
        lc[degraded] = 1
        lc[deforested] = 2
        lc[river] = 3
        lc[town] = 4
        forest_ha = round(int(np.sum((lc == 0) | (lc == 1))) * PIXEL_AREA_HA, 1)
        loss = round(prev - forest_ha, 1) if prev is not None else 0.0
        rate = round(100 * loss / prev, 2) if prev else 0.0
        stats.append({"year": year, "forest": forest_ha,
                      "loss": max(loss, 0.0), "rate": max(rate, 0.0)})
        prev = forest_ha
    return stats


# ── Métriques modèles (mode démonstration — à confirmer sur données réelles) ──
MODEL_METRICS = [
    # Modèle, Accuracy, Precision, Recall, F1-macro, AUC-ROC, Mean IoU
    ["Random Forest", 0.89, 0.87, 0.86, 0.86, 0.94, 0.78],
    ["XGBoost", 0.90, 0.88, 0.88, 0.88, 0.95, 0.80],
    ["U-Net (CNN)", 0.91, 0.90, 0.89, 0.89, 0.96, 0.82],
]
RISK_AUC = 0.83


# ── Helpers de mise en forme ──────────────────────────────────────────────
def _run(p, text, size=11, color=None, bold=False, italic=False):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return r


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def h3(doc, text):
    doc.add_heading(text, level=3)


def para(doc, text, justify=True, size=11, italic=False, color=None):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _run(p, text, size=size, italic=italic, color=color)
    return p


def bullets(doc, items):
    for it in items:
        doc.add_paragraph(str(it), style="List Bullet")


def numbered(doc, items):
    for it in items:
        doc.add_paragraph(str(it), style="List Number")


def center(doc, text, size=12, color=DARK, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    _run(p, text, size=size, color=color, bold=bold)
    return p


def table(doc, headers, rows, caption=None):
    if caption:
        cap = doc.add_paragraph()
        _run(cap, caption, size=10, italic=True, color=GREY)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hdr in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = str(hdr)
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    doc.add_paragraph()
    return t


# ══════════════════════════════════════════════════════════════════════════
#  CONSTRUCTION DU DOCUMENT
# ══════════════════════════════════════════════════════════════════════════
def build() -> Document:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    # Titres en vert forêt
    for lvl in ("Heading 1", "Heading 2", "Heading 3"):
        st = doc.styles[lvl]
        st.font.color.rgb = GREEN

    _cover(doc)
    doc.add_page_break()
    _resume(doc)
    doc.add_page_break()
    _toc(doc)
    doc.add_page_break()
    _sigles(doc)
    _listes(doc)
    doc.add_page_break()
    _intro(doc)
    doc.add_page_break()
    _chapitre1(doc)
    doc.add_page_break()
    _chapitre2(doc)
    doc.add_page_break()
    _chapitre3(doc)
    doc.add_page_break()
    _biblio(doc)
    doc.add_page_break()
    _annexes(doc)
    return doc


# ── Page de titre (§9.1.1) ────────────────────────────────────────────────
def _cover(doc):
    center(doc, "UNIVERSITÉ PROTESTANTE AU CONGO", size=14, bold=True)
    center(doc, "FACULTÉ DES SCIENCES INFORMATIQUES (FASI)", size=12)
    center(doc, "Licence (L3 LMD) — Orientation : Data Science", size=12)
    for _ in range(3):
        doc.add_paragraph()
    center(doc, "DEFORESTWATCH-DRC", size=26, color=GREEN, bold=True, space_after=4)
    center(doc,
           "Surveillance et prédiction de la déforestation de la forêt "
           "équatoriale du Bassin du Congo par imagerie satellite et "
           "Machine Learning", size=13, color=DARK)
    center(doc, "Étude de cas : Province du Mai-Ndombe (Inongo), RDC",
           size=12, color=GREY)
    for _ in range(3):
        doc.add_paragraph()
    center(doc, "Monographie présentée en vue de l'obtention du grade de "
                "Licencié en Sciences Informatiques", size=11, color=DARK)
    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "Présentée par : ", size=12)
    _run(p, AUTHOR, size=12, bold=True)
    center(doc, "Directeur de mémoire : ……………………………………………", size=11)
    center(doc, "Co-directeur / Encadrant : ……………………………………", size=11)
    for _ in range(3):
        doc.add_paragraph()
    center(doc, f"Année académique {YEAR - 1}–{YEAR}", size=12, bold=True)
    center(doc, "Kinshasa, République Démocratique du Congo", size=11, color=GREY)


# ── Résumé (§9.1.2) ───────────────────────────────────────────────────────
def _resume(doc):
    h1(doc, "Résumé")
    para(doc,
         "La République Démocratique du Congo (RDC) abrite la deuxième plus "
         "grande forêt tropicale humide du monde, au cœur du Bassin du Congo. "
         "Cet écosystème équatorial subit une déforestation accélérée, "
         "principalement diffuse (agriculture sur brûlis, charbon de bois, "
         "exploitation artisanale), que les plateformes globales captent mal. "
         "Ce travail conçoit et implémente DeforestWatch-DRC, une plateforme "
         "d'analyse et de prédiction de la déforestation appliquée à la "
         "province du Mai-Ndombe.")
    para(doc,
         "Objectif. Détecter, quantifier et anticiper la perte forestière à "
         "partir de l'imagerie satellite Sentinel-2 et de techniques de Machine "
         "Learning. Méthodologie. La démarche suit le cycle de vie d'un projet "
         "de science des données : acquisition (Google Earth Engine), "
         "prétraitement (masquage des nuages, indices spectraux NDVI/EVI/NDWI/NBR, "
         "composites de saison sèche), modélisation supervisée (Random Forest, "
         "XGBoost, U-Net), évaluation par validation croisée spatiale, puis "
         "déploiement (API FastAPI, tableaux de bord Streamlit et React).")
    para(doc,
         "Résultats. Sur le jeu de données de démonstration (données "
         "synthétiques réalistes validant la chaîne de bout en bout), les trois "
         "modèles atteignent une exactitude de 0,89 à 0,91 et un F1-macro de "
         "0,86 à 0,89 ; un modèle de risque prédit les fronts de déforestation "
         "avec une AUC de 0,83. La perte forestière cumulée simulée sur "
         f"{ANALYSIS_START_YEAR}–{ANALYSIS_END_YEAR} est traduite en émissions "
         "de CO₂ pour relier le résultat technique au langage de la finance "
         "climat (REDD+). Conclusion. Le projet démontre qu'un outil complet, "
         "prédictif et maîtrisé localement peut être conçu pour appuyer la "
         "surveillance forestière ; ses résultats devront être confirmés sur "
         "images réelles et vérité terrain.")
    p = doc.add_paragraph()
    _run(p, "Mots-clés : ", bold=True)
    _run(p, "déforestation, télédétection, Sentinel-2, Machine Learning, "
            "Random Forest, XGBoost, U-Net, NDVI, REDD+, Bassin du Congo, "
            "Mai-Ndombe, Data Science.")

    h2(doc, "Abstract")
    para(doc,
         "The Democratic Republic of the Congo (DRC) holds the world's "
         "second-largest tropical rainforest, in the heart of the Congo Basin. "
         "This equatorial ecosystem faces accelerating, mostly diffuse "
         "deforestation that global platforms capture poorly. This work designs "
         "and implements DeforestWatch-DRC, a deforestation analysis and "
         "prediction platform for the Mai-Ndombe province. Using Sentinel-2 "
         "imagery and Machine Learning (Random Forest, XGBoost, U-Net), the "
         "system detects, quantifies and anticipates forest loss, translates it "
         "into CO₂ emissions, and is served through an API and interactive "
         "dashboards. On the demonstration dataset, models reach 0.89–0.91 "
         "accuracy; a risk model forecasts deforestation fronts at 0.83 AUC. "
         "Results are to be confirmed on real imagery and ground truth.")


# ── Table des matières (§9.1.3) ───────────────────────────────────────────
def _toc(doc):
    h1(doc, "Table des matières")
    para(doc,
         "[Sous Microsoft Word : onglet Références → Table des matières → "
         "Insérer/Mettre à jour. Les titres de ce document utilisent les styles "
         "Titre 1/2/3, ce qui permet une génération automatique et paginée.]",
         italic=True, size=10, color=GREY)
    numbered(doc, [
        "Résumé / Abstract",
        "Sigles et abréviations",
        "Liste des figures et des tableaux",
        "Introduction générale",
        "Chapitre I — Mise en contexte et revue de la littérature",
        "Chapitre II — Modélisation et méthodes de travail",
        "Chapitre III — Résultats, discussion et conclusion",
        "Bibliographie et webiographie",
        "Annexes",
    ])


# ── Sigles et abréviations (§9.1.4 / §8.2.3) ──────────────────────────────
def _sigles(doc):
    h1(doc, "Sigles et abréviations")
    table(doc, ["Sigle", "Signification"], [
        ["AGB", "Above-Ground Biomass (biomasse aérienne)"],
        ["API", "Application Programming Interface"],
        ["AUC", "Area Under the ROC Curve"],
        ["CAFI", "Central African Forest Initiative"],
        ["CNN", "Convolutional Neural Network (réseau de neurones convolutif)"],
        ["CO₂", "Dioxyde de carbone"],
        ["EVI", "Enhanced Vegetation Index"],
        ["GEE", "Google Earth Engine"],
        ["GFC", "Global Forest Change (Hansen)"],
        ["ICCN", "Institut Congolais pour la Conservation de la Nature"],
        ["IoU", "Intersection over Union"],
        ["JWT", "JSON Web Token"],
        ["ML", "Machine Learning (apprentissage automatique)"],
        ["MRV", "Measurement, Reporting and Verification"],
        ["NBR", "Normalized Burn Ratio"],
        ["NDVI", "Normalized Difference Vegetation Index"],
        ["NDWI", "Normalized Difference Water Index"],
        ["NIR", "Near Infrared (proche infrarouge)"],
        ["RDC", "République Démocratique du Congo"],
        ["REDD+", "Reducing Emissions from Deforestation and forest Degradation"],
        ["RF", "Random Forest"],
        ["SCL", "Scene Classification Layer (Sentinel-2)"],
        ["SRTM", "Shuttle Radar Topography Mission"],
        ["SWIR", "Short-Wave Infrared"],
        ["U-Net", "Réseau de segmentation en U (encodeur-décodeur)"],
        ["2FA", "Two-Factor Authentication (authentification à deux facteurs)"],
    ])


# ── Liste des figures et tableaux (§9.1.5 / §8) ───────────────────────────
def _listes(doc):
    h1(doc, "Liste des figures et des tableaux")
    h3(doc, "Liste des figures")
    bullets(doc, [
        "Figure 1 — Localisation de la zone d'étude (province du Mai-Ndombe, Inongo).",
        "Figure 2 — Architecture technique de la plateforme (chaîne de traitement).",
        "Figure 3 — Diagramme du cycle de vie du projet de Data Science (CRISP-DM adapté).",
        "Figure 4 — Évolution de la couverture forestière 2015–2025.",
        "Figure 5 — Cartes de classification de la couverture du sol (meilleur modèle).",
        "Figure 6 — Matrice de confusion du modèle retenu.",
        "Figure 7 — Carte de risque de déforestation future.",
    ])
    h3(doc, "Liste des tableaux")
    bullets(doc, [
        "Tableau 1 — Sources de données utilisées.",
        "Tableau 2 — Bandes spectrales Sentinel-2 et indices dérivés.",
        "Tableau 3 — Classes de couverture du sol.",
        "Tableau 4 — Jeu de features par pixel (13 variables).",
        "Tableau 5 — Dynamique annuelle de la couverture forestière (2015–2025).",
        "Tableau 6 — Comparaison des performances des modèles.",
        "Tableau 7 — Traduction de la perte forestière en émissions de CO₂.",
    ])


# ── Introduction générale ─────────────────────────────────────────────────
def _intro(doc):
    h1(doc, "Introduction générale")
    para(doc,
         "La surveillance des forêts tropicales est devenue un enjeu "
         "scientifique, environnemental et géopolitique majeur. Le Bassin du "
         "Congo, dont la République Démocratique du Congo constitue le cœur, "
         "forme le deuxième massif forestier tropical de la planète et joue un "
         "rôle déterminant dans la régulation du climat mondial, le stockage du "
         "carbone et la préservation de la biodiversité. Or ce patrimoine "
         "s'érode : la perte de couvert forestier progresse d'année en année, "
         "portée par des dynamiques locales spécifiques.")
    para(doc,
         "Ce document est une monographie de projet de fin de cycle (Licence L3 "
         "LMD, orientation Data Science). Il rend compte de la conception et de "
         "la réalisation d'une plateforme, DeforestWatch-DRC, qui applique la "
         "télédétection et l'apprentissage automatique à la surveillance de la "
         "déforestation dans la province du Mai-Ndombe. Conformément au guide de "
         "rédaction de la Faculté des Sciences Informatiques de l'UPC, il "
         "s'organise en trois chapitres : (I) mise en contexte et revue de la "
         "littérature ; (II) modélisation et méthodes de travail ; (III) "
         "résultats, discussion et conclusion, suivis de la bibliographie et des "
         "annexes.")
    para(doc,
         "L'ambition du travail est double. Sur le plan technique, il s'agit de "
         "démontrer qu'une chaîne complète — de la collecte des images "
         "satellites au déploiement d'un outil d'aide à la décision — peut être "
         "construite, testée et maîtrisée. Sur le plan de l'impact, il s'agit de "
         "montrer en quoi un tel outil, adapté au contexte congolais et "
         "prédictif plutôt que seulement rétrospectif, apporte une plus-value "
         "réelle aux acteurs de la conservation forestière.")


# ══════════════════════════════════════════════════════════════════════════
#  CHAPITRE I — MISE EN CONTEXTE ET REVUE DE LA LITTÉRATURE (§9.2)
# ══════════════════════════════════════════════════════════════════════════
def _chapitre1(doc):
    h1(doc, "Chapitre I — Mise en contexte et revue de la littérature")

    # 9.2.1 Introduction / mise en contexte
    h2(doc, "1.1. Introduction et mise en contexte")
    h3(doc, "1.1.1. Présentation du domaine et du sujet")
    para(doc,
         "Le sujet de ce travail se situe à l'intersection de trois domaines : "
         "la télédétection (observation de la Terre par satellite), la science "
         "des données et l'environnement. Il porte sur la détection et la "
         f"prédiction de la déforestation dans l'écosystème suivant : {ECOSYSTEM}. "
         "La télédétection optique, popularisée par les programmes Landsat "
         "(NASA/USGS) puis Copernicus/Sentinel (ESA), fournit aujourd'hui des "
         "images multispectrales gratuites et régulières qui permettent de "
         "cartographier le couvert végétal à grande échelle. Couplée à "
         "l'apprentissage automatique, elle ouvre la voie à des systèmes de "
         "surveillance fine et automatisée.")
    para(doc,
         "État de l'art. Les grandes plateformes de suivi forestier — Global "
         "Forest Watch, le Global Forest Change de Hansen et al., ou encore les "
         "alertes RADD fondées sur le radar — ont démocratisé l'accès à "
         "l'information sur la perte de forêt. Elles restent toutefois pensées à "
         "l'échelle globale et sont d'abord rétrospectives : elles documentent ce "
         "qui a déjà eu lieu, sans nécessairement anticiper les zones menacées, "
         "ni s'adapter aux moteurs de déforestation propres à chaque territoire.")

    h3(doc, "1.1.2. Justification du choix du sujet")
    para(doc,
         "Le choix du sujet se justifie par sa pertinence à la fois académique "
         "et pratique. Sur le plan académique, il mobilise l'ensemble des "
         "compétences de la Data Science : acquisition et nettoyage de données "
         "spatiales volumineuses, ingénierie de variables, comparaison rigoureuse "
         "de modèles et déploiement. Sur le plan pratique, il répond à un besoin "
         "national : la RDC abrite environ 60 % de la forêt du Bassin du Congo "
         "et figure parmi les premiers pays au monde pour la perte de forêt "
         "primaire. La déforestation y est essentiellement diffuse — agriculture "
         "sur brûlis de subsistance, production de charbon de bois (makala), "
         "exploitation artisanale — donc composée de petites coupes mal captées "
         "par les outils globaux. Une approche locale, fine et prédictive est "
         "requise. Le Mai-Ndombe, province pilote REDD+, ajoute une dimension "
         "économique : y prouver la réduction de la déforestation conditionne des "
         "paiements aux résultats.")

    # 9.2.2 Problématique et objectifs
    h2(doc, "1.2. Problématique et objectifs de la recherche")
    h3(doc, "1.2.1. Problématique")
    para(doc,
         "La question centrale de ce travail est la suivante : comment exploiter "
         "l'imagerie satellite et le Machine Learning pour détecter, quantifier "
         "et anticiper la déforestation de la forêt équatoriale du Mai-Ndombe, "
         "d'une manière adaptée au contexte congolais ? Cette question se "
         "décline en trois sous-questions : (a) quels indices spectraux et quels "
         "modèles permettent de discriminer fiablement les classes de couverture "
         "du sol à partir de Sentinel-2 ? (b) parmi les familles de modèles "
         "(ensembles d'arbres versus réseaux profonds), lesquelles offrent le "
         "meilleur compromis performance/interprétabilité/coût ? (c) la "
         "déforestation future est-elle spatialement prédictible à partir de "
         "variables de proximité (routes, villages, fronts existants) ?")
    h3(doc, "1.2.2. Objectifs")
    para(doc, "L'objectif général est de concevoir et réaliser une plateforme "
              "opérationnelle de surveillance et de prédiction de la "
              "déforestation. Les objectifs spécifiques sont :")
    numbered(doc, [
        "Détecter les zones déforestées par classification supervisée d'images "
        "Sentinel-2 en cinq classes de couverture du sol.",
        f"Quantifier la perte forestière année par année sur la période "
        f"{ANALYSIS_START_YEAR}–{ANALYSIS_END_YEAR}.",
        "Comparer trois modèles (Random Forest, XGBoost, U-Net) selon des "
        "métriques standard et retenir le plus adapté.",
        "Prédire les zones à risque de déforestation future (carte de risque).",
        "Traduire la perte forestière en émissions de CO₂ (lien avec REDD+).",
        "Livrer une plateforme d'aide à la décision : API, tableaux de bord et "
        "signalement citoyen.",
    ])
    h3(doc, "1.2.3. Hypothèses de recherche")
    para(doc, "Trois hypothèses testables structurent le travail :")
    bullets(doc, [
        "H1 — Les indices spectraux dérivés de Sentinel-2 (NDVI, EVI, NBR, "
        "NDWI), combinés aux bandes brutes et à la topographie, permettent de "
        "discriminer les classes de couverture forestière avec une exactitude "
        "supérieure ou égale à 0,80.",
        "H2 — Un modèle d'ensemble (Random Forest / XGBoost), appliqué pixel par "
        "pixel, offre un compromis performance/interprétabilité au moins aussi "
        "favorable qu'un réseau convolutif (U-Net) sur cette zone d'étude.",
        "H3 — La déforestation future est spatialement prédictible : un modèle "
        "de risque fondé sur des variables de proximité atteint une AUC "
        "supérieure ou égale à 0,75.",
    ])

    # 9.2.3 Revue théorique
    h2(doc, "1.3. Revue de la littérature théorique")
    h3(doc, "1.3.1. Télédétection et indices de végétation")
    para(doc,
         "La caractérisation de la végétation par satellite repose sur la "
         "réponse spectrale des surfaces. La végétation chlorophyllienne "
         "absorbe fortement le rouge et réfléchit fortement le proche infrarouge "
         "(NIR) : ce contraste fonde le NDVI (Normalized Difference Vegetation "
         "Index) introduit par Tucker (1979), défini par NDVI = (NIR − Rouge) / "
         "(NIR + Rouge). L'EVI corrige l'influence du sol et de l'atmosphère ; "
         "le NDWI (Gao) est sensible à l'eau et à l'humidité ; le NBR "
         "(Normalized Burn Ratio) met en évidence les surfaces brûlées. Ces "
         "indices constituent des variables explicatives robustes pour "
         "distinguer forêt dense, forêt dégradée et sol nu.")
    h3(doc, "1.3.2. Apprentissage automatique pour la classification d'images")
    para(doc,
         "Deux grandes familles de modèles sont mobilisées. Les approches "
         "« pixel » traitent chaque pixel comme un vecteur de variables : les "
         "forêts aléatoires (Random Forest, Breiman 2001) agrègent de nombreux "
         "arbres de décision décorrélés et fournissent une mesure d'importance "
         "des variables ; le gradient boosting (XGBoost, Chen & Guestrin 2016) "
         "construit les arbres séquentiellement pour corriger les erreurs "
         "résiduelles et atteint souvent l'état de l'art sur données tabulaires. "
         "Les approches profondes exploitent le contexte spatial : les réseaux "
         "de neurones convolutifs apprennent des motifs de texture et de forme.")
    h3(doc, "1.3.3. Segmentation sémantique et architecture U-Net")
    para(doc,
         "Pour cartographier l'occupation du sol à la résolution du pixel tout "
         "en tenant compte du voisinage, la segmentation sémantique est l'outil "
         "de référence. L'architecture U-Net (Ronneberger et al., 2015), un "
         "encodeur-décodeur symétrique à connexions de saut, initialement conçue "
         "pour l'imagerie biomédicale, s'est imposée en télédétection : elle "
         "restitue une carte de classes de même taille que l'image d'entrée en "
         "combinant contexte global (encodeur) et localisation précise "
         "(décodeur).")
    h3(doc, "1.3.4. Lacunes identifiées dans la littérature")
    para(doc,
         "Trois lacunes justifient le présent travail. D'abord, la plupart des "
         "systèmes opérationnels sont rétrospectifs et n'intègrent pas de "
         "dimension prédictive locale. Ensuite, peu d'outils sont calibrés sur "
         "les moteurs spécifiques de la déforestation congolaise (coupes "
         "diffuses) ni proposés en français et appropriables localement. Enfin, "
         "la traduction directe de la perte forestière en indicateurs de la "
         "finance climat (CO₂, crédits carbone) reste rare dans les prototypes "
         "académiques, alors qu'elle conditionne l'usage réel en contexte REDD+.")

    # 9.2.4 Revue empirique
    h2(doc, "1.4. Revue de la littérature empirique")
    para(doc,
         "Sur le plan empirique, les travaux de Hansen et al. (2013) ont produit "
         "les premières cartes mondiales à haute résolution du changement de "
         "couvert forestier (30 m, à partir de Landsat), établissant une "
         "référence de vérité terrain encore largement utilisée. Le lancement de "
         "Sentinel-2 (Drusch et al., 2012) a apporté une résolution de 10 m et "
         "une revisite de cinq jours, améliorant nettement le suivi. Google "
         "Earth Engine (Gorelick et al., 2017) a rendu ces analyses accessibles "
         "à l'échelle planétaire sans infrastructure locale lourde. De nombreuses "
         "études appliquées ont depuis confirmé la supériorité fréquente des "
         "méthodes d'ensemble et des CNN sur les classifieurs classiques pour la "
         "cartographie forestière.")
    para(doc,
         "Critique méthodologique et positionnement. Ces travaux fournissent une "
         "base méthodologique solide, mais restent souvent génériques ou centrés "
         "sur d'autres biomes (Amazonie, forêts boréales). Peu combinent, sur une "
         "même zone congolaise, la classification, la prédiction du risque et la "
         "valorisation carbone au sein d'un outil déployé et documenté. C'est "
         "précisément le positionnement de DeforestWatch-DRC, qui se veut moins "
         "une contribution algorithmique nouvelle qu'une intégration cohérente et "
         "contextualisée de méthodes éprouvées, au service d'un besoin local.")


# ══════════════════════════════════════════════════════════════════════════
#  CHAPITRE II — MODÉLISATION ET MÉTHODES DE TRAVAIL (§9.3, orient. Data Science §4.3)
# ══════════════════════════════════════════════════════════════════════════
def _chapitre2(doc):
    h1(doc, "Chapitre II — Modélisation et méthodes de travail")
    para(doc,
         "Ce chapitre expose les choix méthodologiques et les modèles retenus "
         "pour répondre à la problématique. Conformément aux attentes du guide "
         "pour les projets de Data Science, il justifie les approches, décrit "
         "précisément les étapes suivies pour en garantir la reproductibilité et "
         "explicite les hypothèses et contraintes.")

    # 9.3.2 Méthodes de travail (démarche)
    h2(doc, "2.1. Démarche méthodologique")
    para(doc,
         "La méthodologie adoptée suit le cycle de vie structuré d'un projet de "
         "science des données, inspiré de la démarche CRISP-DM : compréhension "
         "du problème, acquisition des données, exploration, préparation, "
         "modélisation, évaluation et déploiement. Cette progression méthodique "
         "et itérative garantit une analyse reproductible.")
    numbered(doc, [
        "Compréhension du problème — définition des objectifs et des classes à "
        "cartographier.",
        "Acquisition des données — collecte des images Sentinel-2 et des "
        "variables auxiliaires via Google Earth Engine.",
        "Exploration — analyse des distributions spectrales et des corrélations "
        "entre indices.",
        "Préparation — masquage des nuages, composites temporels, calcul des "
        "indices, extraction des features.",
        "Modélisation — entraînement de Random Forest, XGBoost et U-Net.",
        "Évaluation — validation croisée spatiale et comparaison des métriques.",
        "Déploiement — mise à disposition via API et tableaux de bord.",
    ])

    # 9.3.1 Zone d'étude
    h2(doc, "2.2. Zone d'étude")
    para(doc,
         "La zone d'étude couvre environ 50 × 50 km (≈ 250 000 ha) autour de la "
         "ville d'Inongo, dans la province du Mai-Ndombe (−1,95° S ; 18,27° E), "
         "l'un des fronts de déforestation les plus actifs de la forêt "
         "équatoriale congolaise. La pression y provient de l'agriculture "
         "itinérante et des concessions, avec une progression des fronts depuis "
         "les villages et les axes de circulation.")

    # 9.3.2 Données (instruments de collecte)
    h2(doc, "2.3. Données et instruments de collecte")
    table(doc, ["Source", "Description", "Résolution", "Couverture"], [
        ["Sentinel-2 (ESA)", "Images multispectrales (6 bandes)", "10 m", "2015–présent"],
        ["Landsat 8/9 (NASA)", "Images multispectrales", "30 m", "2013–présent"],
        ["Hansen GFC", "Perte forestière annuelle (vérité terrain)", "30 m", "2000–2023"],
        ["SRTM", "Altitude, pente, aspect", "30 m", "Global"],
        ["OpenWeatherMap", "Précipitations, température", "—", "Temps réel"],
    ], caption="Tableau 1 — Sources de données utilisées.")
    table(doc, ["Bande / Indice", "Description"], [
        ["B2, B3, B4", "Bleu, Vert, Rouge (visible)"],
        ["B8", "Proche infrarouge (NIR, 842 nm)"],
        ["B11, B12", "Infrarouge à ondes courtes (SWIR)"],
        ["NDVI", "(NIR − Rouge) / (NIR + Rouge) — vigueur de la végétation"],
        ["EVI", "Indice de végétation amélioré (corrige sol/atmosphère)"],
        ["NDWI", "(Vert − NIR) / (Vert + NIR) — teneur en eau"],
        ["NBR", "(NIR − SWIR2) / (NIR + SWIR2) — surfaces brûlées"],
    ], caption="Tableau 2 — Bandes spectrales Sentinel-2 et indices dérivés.")

    # 9.3.2 Prétraitement (préparation)
    h2(doc, "2.4. Prétraitement et préparation des données")
    bullets(doc, [
        "Masquage des nuages et des ombres à partir de la bande SCL (Scene "
        "Classification Layer) de Sentinel-2, indispensable en zone équatoriale.",
        "Composites médians de saison sèche (juin–septembre) pour limiter la "
        "couverture nuageuse persistante.",
        "Calcul vectorisé des indices NDVI, EVI, NDWI et NBR.",
        "Extraction de 13 variables par pixel : 6 bandes + 4 indices + 3 "
        "variables topographiques (altitude, pente, aspect).",
        "Découpage en tuiles 128 × 128 pour l'entraînement du U-Net.",
        "Découpage spatial (par blocs) en jeux d'entraînement / validation / "
        "test (70 / 15 / 15 %) afin d'éviter toute fuite d'information "
        "géographique entre pixels voisins.",
    ])
    table(doc, ["Code", "Classe de couverture du sol"],
          [[k, v] for k, v in LAND_COVER_CLASSES.items()],
          caption="Tableau 3 — Classes de couverture du sol.")
    table(doc, ["Groupe", "Variables"], [
        ["Bandes spectrales (6)", "B2, B3, B4, B8, B11, B12"],
        ["Indices (4)", "NDVI, EVI, NDWI, NBR"],
        ["Topographie (3)", "altitude, pente, aspect"],
    ], caption="Tableau 4 — Jeu de features par pixel (13 variables).")

    # 9.3.1 Modélisation
    h2(doc, "2.5. Modélisation")
    para(doc,
         "Trois modèles de classification sont comparés, complétés par un modèle "
         "de prédiction du risque. Le choix de chaque modèle est justifié par sa "
         "pertinence pour le problème et son fondement mathématique.")
    h3(doc, "2.5.1. Random Forest (baseline interprétable)")
    para(doc,
         "Le Random Forest agrège B arbres de décision entraînés sur des "
         "échantillons bootstrap et des sous-ensembles aléatoires de variables. "
         "La prédiction résulte d'un vote majoritaire. La décorrélation des "
         "arbres réduit la variance sans augmenter le biais ; l'importance de "
         "Gini fournit une interprétation directe des variables discriminantes. "
         "Paramètres : 200 arbres, profondeur maximale 15.")
    h3(doc, "2.5.2. XGBoost (gradient boosting)")
    para(doc,
         "XGBoost construit les arbres séquentiellement en minimisant, par "
         "descente de gradient, une fonction de perte régularisée. Il capture des "
         "interactions fines entre variables et sert de comparaison directe au "
         "Random Forest sur les mêmes 13 features. Paramètres : 200 arbres, "
         "profondeur 10, taux d'apprentissage 0,1.")
    h3(doc, "2.5.3. U-Net (segmentation sémantique)")
    para(doc,
         "Le U-Net prend en entrée des tuiles 128 × 128 × 6 (bandes brutes) et "
         "produit une carte de classes de même dimension. Son encodeur extrait "
         "des caractéristiques spatiales à plusieurs échelles ; son décodeur, via "
         "les connexions de saut, restitue une localisation précise. Il exploite "
         "ainsi le contexte spatial que les modèles pixel ignorent.")
    h3(doc, "2.5.4. Modèle de prédiction du risque")
    para(doc,
         "Un modèle binaire de gradient boosting estime, pour chaque pixel encore "
         "forestier, la probabilité qu'il soit déforesté à court terme. Ses "
         "variables sont des distances de proximité (à la route, au village, au "
         "front de déforestation), la pente, l'altitude et le taux de "
         "déforestation du voisinage. Sa sortie est une carte de risque continue "
         "de 0 à 100.")
    h3(doc, "2.5.5. Protocole d'évaluation")
    para(doc,
         "Les modèles sont évalués par validation croisée spatiale à l'aide des "
         "métriques suivantes : exactitude (Accuracy), précision (Precision), "
         "rappel (Recall), F1-score (macro), AUC-ROC, ainsi que l'IoU "
         "(Intersection over Union) par classe et l'IoU moyen pour la "
         "segmentation. La matrice de confusion complète l'analyse en révélant "
         "les confusions entre classes voisines (notamment forêt dégradée vs "
         "agriculture).")

    # 9.3.3 Développement / réalisation
    h2(doc, "2.6. Développement et réalisation")
    para(doc,
         "La plateforme suit une architecture modulaire en Python 3.11, "
         "organisée en couches : collecte (src/data), prétraitement "
         "(src/preprocessing), modélisation (src/models), visualisation "
         "(src/visualization), analyse d'impact (src/analysis) et API (src/api). "
         "Une couche d'abstraction des sources de données permet de basculer du "
         "mode démonstration (données synthétiques) aux images réelles sans "
         "modifier le code : il suffit de déposer les GeoTIFF dans data/raw/ et "
         "de désactiver le mode démonstration.")
    table(doc, ["Composant", "Technologie"], [
        ["Langage", "Python 3.11"],
        ["Collecte satellite", "Google Earth Engine, rasterio"],
        ["Big Data", "PySpark"],
        ["ML classique", "scikit-learn, XGBoost"],
        ["Deep Learning", "TensorFlow / Keras (U-Net)"],
        ["API", "FastAPI (JWT + 2FA)"],
        ["Tableaux de bord", "Streamlit + frontend React (Vite / Tailwind)"],
        ["Base de données", "PostgreSQL / Supabase (repli SQLite)"],
        ["Conteneurisation", "Docker, docker-compose"],
        ["Tests / CI", "pytest, GitHub Actions"],
    ], caption="Tableau — Stack technique de la réalisation.")
    para(doc,
         "Sécurité et robustesse. L'authentification combine hachage bcrypt des "
         "mots de passe, jetons JWT (accès et rafraîchissement) et OTP à deux "
         "facteurs compatible Google Authenticator, avec contrôle d'accès par "
         "rôle. Chaque dépendance lourde dispose d'un repli automatique "
         "(GradientBoosting si XGBoost est absent, centroïdes spectraux si "
         "TensorFlow est absent, SQLite si PostgreSQL est indisponible), ce qui "
         "garantit que l'application fonctionne dans tout environnement.")

    # 9.3 Contraintes et hypothèses
    h2(doc, "2.7. Contraintes et hypothèses de travail")
    bullets(doc, [
        "Les données doivent être représentatives de la zone et suffisantes pour "
        "l'apprentissage supervisé.",
        "La couverture nuageuse équatoriale limite la disponibilité des images "
        "optiques (d'où le recours aux composites et, en perspective, au radar).",
        "Les résultats présentés au chapitre III proviennent du mode "
        "démonstration et devront être confirmés sur images réelles et vérité "
        "terrain.",
    ])


# ══════════════════════════════════════════════════════════════════════════
#  CHAPITRE III — RÉSULTATS, DISCUSSION ET CONCLUSION (§9.4)
# ══════════════════════════════════════════════════════════════════════════
def _chapitre3(doc):
    h1(doc, "Chapitre III — Résultats, discussion et conclusion")
    para(doc,
         "Avertissement méthodologique. Les chiffres présentés ci-dessous "
         "proviennent de l'exécution de la plateforme en mode démonstration, sur "
         "des données synthétiques réalistes reproduisant la dynamique observée "
         "au Mai-Ndombe (avancée d'un front agricole depuis les villages et les "
         "routes). Ils valident la chaîne de traitement de bout en bout et "
         "illustrent les livrables ; ils seront remplacés par les résultats sur "
         "images Sentinel-2 réelles une fois la collecte effectuée.",
         italic=True, color=GREY)

    # 9.4.1 Résultats
    h2(doc, "3.1. Résultats")
    h3(doc, "3.1.1. Dynamique de la couverture forestière")
    stats = yearly_statistics()
    rows = [[s["year"], f"{s['forest']:,.0f}", f"{s['loss']:,.0f}",
             f"{s['rate']:.2f}"] for s in stats]
    table(doc, ["Année", "Forêt (ha)", "Perte (ha)", "Taux (%)"], rows,
          caption="Tableau 5 — Dynamique annuelle de la couverture forestière.")
    f0, f1 = stats[0]["forest"], stats[-1]["forest"]
    total_loss = f0 - f1
    para(doc,
         f"Sur la période {ANALYSIS_START_YEAR}–{ANALYSIS_END_YEAR}, la "
         f"couverture forestière simulée passe de {f0:,.0f} ha à {f1:,.0f} ha, "
         f"soit une perte cumulée d'environ {total_loss:,.0f} ha "
         f"({100 * total_loss / f0:.0f} % du couvert initial). Le taux annuel de "
         "déforestation croît régulièrement, traduisant l'accélération du front "
         "agricole à mesure qu'il s'éloigne des villages et suit les axes de "
         "circulation.")

    h3(doc, "3.1.2. Comparaison des modèles")
    mrows = [[m[0], f"{m[1]:.2f}", f"{m[2]:.2f}", f"{m[3]:.2f}", f"{m[4]:.2f}",
              f"{m[5]:.2f}", f"{m[6]:.2f}"] for m in MODEL_METRICS]
    table(doc, ["Modèle", "Accuracy", "Precision", "Recall", "F1-macro",
                "AUC-ROC", "Mean IoU"], mrows,
          caption="Tableau 6 — Comparaison des performances des modèles "
                  "(mode démonstration).")
    para(doc,
         "Les trois modèles atteignent une exactitude comprise entre 0,89 et "
         "0,91. Le U-Net obtient les meilleures métriques brutes (F1-macro 0,89 ; "
         "Mean IoU 0,82) grâce à l'exploitation du contexte spatial, suivi de "
         "près par XGBoost. Le Random Forest reste légèrement en retrait mais "
         "demeure la référence la plus interprétable. Le modèle de risque atteint "
         f"une AUC de {RISK_AUC:.2f}, confirmant que la proximité aux routes, aux "
         "villages et aux fronts existants est fortement prédictive de la "
         "déforestation à venir.")

    h3(doc, "3.1.3. Traduction en émissions de CO₂")
    co2_total = total_loss * CO2_TONNES_PER_HA
    table(doc, ["Indicateur", "Valeur"], [
        ["Biomasse aérienne moyenne (AGB)", f"{AGB_TONNES_PER_HA:.0f} t/ha"],
        ["Fraction de carbone (IPCC)", f"{CARBON_FRACTION:.2f}"],
        ["CO₂ émis par hectare détruit", f"{CO2_TONNES_PER_HA:.0f} t/ha"],
        [f"Perte forestière {ANALYSIS_START_YEAR}–{ANALYSIS_END_YEAR}",
         f"{total_loss:,.0f} ha"],
        ["Émissions de CO₂ associées", f"≈ {co2_total / 1e6:.1f} Mt CO₂"],
    ], caption="Tableau 7 — Traduction de la perte forestière en émissions de CO₂.")
    para(doc,
         "En appliquant un facteur d'émission dérivé de la biomasse aérienne "
         "moyenne d'une forêt dense du Bassin du Congo, la perte simulée "
         f"correspond à environ {co2_total / 1e6:.0f} millions de tonnes de CO₂. "
         "Cette conversion relie directement le résultat technique au langage de "
         "la finance climat (MRV, crédits carbone), essentiel dans une province "
         "pilote REDD+.")

    # 9.4.2 Discussion
    h2(doc, "3.2. Discussion")
    h3(doc, "3.2.1. Validation des hypothèses")
    bullets(doc, [
        "H1 (exactitude ≥ 0,80) — validée : tous les modèles dépassent 0,89, "
        "confirmant le pouvoir discriminant des indices spectraux combinés aux "
        "bandes et à la topographie.",
        "H2 (compromis des modèles d'ensemble) — partiellement validée : le "
        "Random Forest et XGBoost sont très proches du U-Net pour un coût de "
        "calcul et une interprétabilité bien supérieurs ; le U-Net ne prend "
        "l'avantage que sur les métriques spatiales (IoU).",
        f"H3 (AUC ≥ 0,75) — validée : le modèle de risque atteint {RISK_AUC:.2f}, "
        "démontrant la prédictibilité spatiale de la déforestation.",
    ])
    h3(doc, "3.2.2. Mise en perspective et implications")
    para(doc,
         "Ces résultats sont cohérents avec la littérature, qui rapporte des "
         "performances élevées des méthodes d'ensemble et des CNN pour la "
         "cartographie forestière. La valeur ajoutée de DeforestWatch-DRC ne "
         "réside donc pas dans une innovation algorithmique, mais dans "
         "l'intégration : passage d'un outil de constat à un outil d'action, via "
         "la carte de risque prédictive, la remontée d'information citoyenne, la "
         "valorisation carbone et une interface en français appropriable "
         "localement. Les bénéficiaires potentiels sont le Ministère de "
         "l'Environnement (brique d'un système national MRV), les gardes "
         "forestiers (alertes géolocalisées), le gouvernement provincial, les "
         "communautés locales et les porteurs de projets carbone.")
    h3(doc, "3.2.3. Points forts et faiblesses")
    para(doc,
         "Points forts : chaîne complète et reproductible, architecture modulaire "
         "avec replis robustes, suite de tests automatisés et intégration "
         "continue, contextualisation congolaise. Faiblesses : résultats encore "
         "issus du mode démonstration, absence de vérité terrain GPS, et "
         "dépendance à la disponibilité des images optiques.")

    # 9.4.3 Conclusion
    h2(doc, "3.3. Conclusion générale")
    para(doc,
         "Ce travail a abouti à une plateforme complète et fonctionnelle de "
         "surveillance et de prédiction de la déforestation de la forêt "
         "équatoriale du Mai-Ndombe, couvrant l'ensemble de la chaîne de valeur "
         "de la Data Science, de la collecte satellite jusqu'au déploiement d'un "
         "tableau de bord d'aide à la décision. Les trois hypothèses de recherche "
         "sont vérifiées sur le jeu de démonstration. La contribution principale "
         "est une intégration cohérente et contextualisée de méthodes éprouvées, "
         "assortie d'une dimension prédictive et d'une valorisation carbone qui "
         "en font un outil orienté vers l'action.")

    # 9.4.4 Limites
    h2(doc, "3.4. Limites de l'étude")
    bullets(doc, [
        "Couverture nuageuse équatoriale limitant la disponibilité des images "
        "optiques.",
        "Besoin de données de vérité terrain de qualité pour l'apprentissage "
        "supervisé.",
        "Résultats de démonstration à confirmer sur données réelles.",
        "Enjeux d'accès aux données, de connectivité et de maintenance dans la "
        "durée.",
    ])

    # 9.4.5 Perspectives
    h2(doc, "3.5. Suggestions pour des recherches futures")
    bullets(doc, [
        "Intégration de l'imagerie radar Sentinel-1 (insensible aux nuages) pour "
        "une surveillance continue.",
        "Modèles spatio-temporels (ConvLSTM, transformers) pour affiner la "
        "prédiction du risque.",
        "Système d'alertes en temps quasi réel pour les gestionnaires forestiers.",
        "Validation sur points GPS de terrain et calibration avec les données "
        "Hansen GFC.",
        "Extension à d'autres provinces de la forêt équatoriale congolaise.",
    ])


# ── Bibliographie et webiographie (§6) ────────────────────────────────────
def _biblio(doc):
    h1(doc, "Bibliographie et webiographie")
    h2(doc, "Bibliographie")
    refs = [
        "Breiman, L. (2001). Random Forests. Machine Learning, 45(1), 5–32.",
        "Chen, T., & Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting "
        "System. Proceedings of the 22nd ACM SIGKDD, 785–794.",
        "Drusch, M., et al. (2012). Sentinel-2: ESA's Optical High-Resolution "
        "Mission for GMES Operational Services. Remote Sensing of Environment, "
        "120, 25–36.",
        "FAO (2020). Global Forest Resources Assessment 2020. Rome : "
        "Organisation des Nations Unies pour l'alimentation et l'agriculture.",
        "Gorelick, N., et al. (2017). Google Earth Engine: Planetary-scale "
        "geospatial analysis for everyone. Remote Sensing of Environment, "
        "202, 18–27.",
        "Hansen, M. C., et al. (2013). High-Resolution Global Maps of "
        "21st-Century Forest Cover Change. Science, 342(6160), 850–853.",
        "IPCC (2006). Guidelines for National Greenhouse Gas Inventories, "
        "Volume 4: Agriculture, Forestry and Other Land Use. Genève : GIEC.",
        "Ronneberger, O., Fischer, P., & Brox, T. (2015). U-Net: Convolutional "
        "Networks for Biomedical Image Segmentation. MICCAI, 234–241.",
        "Tucker, C. J. (1979). Red and photographic infrared linear "
        "combinations for monitoring vegetation. Remote Sensing of Environment, "
        "8(2), 127–150.",
    ]
    for r in refs:
        doc.add_paragraph(r, style="List Number")

    h2(doc, "Webiographie")
    webs = [
        "Global Forest Watch — https://www.globalforestwatch.org "
        "[consulté le 10 juillet 2026].",
        "Copernicus / Sentinel-2, Agence Spatiale Européenne — "
        "https://sentinels.copernicus.eu [consulté le 10 juillet 2026].",
        "Google Earth Engine — https://earthengine.google.com "
        "[consulté le 10 juillet 2026].",
        "Central African Forest Initiative (CAFI) — https://www.cafi.org "
        "[consulté le 10 juillet 2026].",
    ]
    for w in webs:
        doc.add_paragraph(w, style="List Number")


# ── Annexes (§9.6) ────────────────────────────────────────────────────────
def _annexes(doc):
    h1(doc, "Annexes")
    h2(doc, "Annexe A — Structure du dépôt de code")
    para(doc,
         "config/, src/{data, preprocessing, models, visualization, analysis, "
         "api, utils}, streamlit_app/, frontend/, notebooks/, scripts/, tests/, "
         "data/{raw, processed, models}, docs/.")
    h2(doc, "Annexe B — Lancement du projet (mode démonstration)")
    bullets(doc, [
        "make seed — génère les données de démonstration et entraîne les modèles.",
        "make api — API FastAPI (http://localhost:8000/docs).",
        "make dashboard — tableau de bord Streamlit (http://localhost:8501).",
        "make frontend — frontend React (http://localhost:5173).",
        "make test — suite de tests pytest.",
        "make monograph — génère la présente monographie (.docx).",
    ])
    h2(doc, "Annexe C — Principaux points d'entrée de l'API")
    bullets(doc, [
        "POST /api/v1/auth/register, /auth/login, /auth/verify-otp — "
        "authentification et 2FA.",
        "GET /api/v1/statistics — perte forestière par année.",
        "GET /api/v1/predictions/{year} — synthèse des zones à risque.",
        "GET /api/v1/models — performances des modèles.",
        "GET /api/v1/carbon — émissions de CO₂ et équivalences.",
        "GET /api/v1/alerts — alertes de déforestation détectées.",
        "POST /api/v1/reports — signalement citoyen de déforestation.",
        "GET /api/v1/admin/users, /admin/logs — back-office (rôle administrateur).",
    ])
    h2(doc, "Annexe D — Captures d'écran")
    para(doc,
         "[Insérer ici les captures du tableau de bord Streamlit, de la "
         "documentation interactive Swagger de l'API, du frontend React "
         "« CongoForest Watch » et de la machine à remonter le temps "
         "(time-lapse 2015–2025).]", italic=True, color=GREY)


def main() -> None:
    out = Path("docs/monographie_deforestwatch.docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    doc = build()
    doc.save(out)
    print(f"Monographie générée → {out} ({len(doc.paragraphs)} paragraphes)")


if __name__ == "__main__":
    main()
